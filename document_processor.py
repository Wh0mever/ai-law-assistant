import os
import logging
from typing import Optional
import docx
import PyPDF2
import fitz  # PyMuPDF
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
    
    def extract_text(self, file_path: str) -> Optional[str]:
        """Извлечение текста из файла"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """Извлечение текста из PDF"""
        try:
            # Пробуем с PyMuPDF (более надежный)
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text += page.get_text()
            
            doc.close()
            
            if text.strip():
                return text
            
            # Если PyMuPDF не дал результата, пробуем PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page in pdf_reader.pages:
                    text += page.extract_text()
                
                return text if text.strip() else None
                
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return None
    
    def _extract_from_docx(self, file_path: str) -> Optional[str]:
        """Извлечение текста из DOCX"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip() if text.strip() else None
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return None
    
    def _extract_from_txt(self, file_path: str) -> Optional[str]:
        """Извлечение текста из TXT"""
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'cp866', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        return text.strip() if text.strip() else None
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Could not decode text file with any encoding")
            return None
            
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            return None
    
    def validate_file(self, file_path: str, max_size_mb: int = 10) -> tuple[bool, str]:
        """Валидация файла"""
        try:
            # Проверяем существование файла
            if not os.path.exists(file_path):
                return False, "Файл не найден"
            
            # Проверяем размер файла
            file_size = os.path.getsize(file_path)
            if file_size > max_size_mb * 1024 * 1024:
                return False, f"Размер файла превышает {max_size_mb} МБ"
            
            # Проверяем формат файла
            file_extension = Path(file_path).suffix.lower().replace('.', '')
            if file_extension not in self.supported_formats:
                return False, f"Поддерживаются только форматы: {', '.join(self.supported_formats)}"
            
            return True, "OK"
            
        except Exception as e:
            logger.error(f"Error validating file {file_path}: {e}")
            return False, "Ошибка при проверке файла"
    
    def clean_text(self, text: str) -> str:
        """Очистка и нормализация текста"""
        if not text:
            return ""
        
        # Удаляем лишние пробелы и переносы строк
        text = ' '.join(text.split())
        
        # Удаляем специальные символы, если нужно
        # text = re.sub(r'[^\w\s\-.,;:!?()«»""\']+', '', text)
        
        return text.strip() 